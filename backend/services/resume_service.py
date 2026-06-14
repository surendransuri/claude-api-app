import os
import uuid
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib import colors

TEMP_DIR = Path("temp_files")
TEMP_DIR.mkdir(exist_ok=True)


def _parse_markdown_to_sections(md_text: str) -> list:
    """Parse markdown into a list of (type, content) tuples."""
    sections = []
    for line in md_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            sections.append(("h1", stripped[2:]))
        elif stripped.startswith("## "):
            sections.append(("h2", stripped[3:]))
        elif stripped.startswith("### "):
            sections.append(("h3", stripped[4:]))
        elif stripped.startswith("**") and stripped.endswith("**"):
            sections.append(("bold", stripped[2:-2]))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            sections.append(("bullet", stripped[2:]))
        elif stripped == "" or stripped == "---":
            sections.append(("separator", ""))
        else:
            sections.append(("text", stripped))
    return sections


def generate_docx(markdown_content: str) -> str:
    doc = Document()

    # Default styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    sections = _parse_markdown_to_sections(markdown_content)

    for stype, content in sections:
        if not content and stype == "separator":
            doc.add_paragraph()
            continue
        if not content:
            continue

        if stype == "h1":
            p = doc.add_heading(content, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)
        elif stype == "h2":
            p = doc.add_heading(content, level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif stype == "h3":
            doc.add_heading(content, level=3)
        elif stype == "bold":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
        elif stype == "bullet":
            doc.add_paragraph(content, style="List Bullet")
        else:
            doc.add_paragraph(content)

    filename = f"resume_{uuid.uuid4().hex[:8]}.docx"
    filepath = TEMP_DIR / filename
    doc.save(str(filepath))
    return filename


def generate_pdf(markdown_content: str) -> str:
    filename = f"resume_{uuid.uuid4().hex[:8]}.pdf"
    filepath = TEMP_DIR / filename

    doc = SimpleDocTemplate(
        str(filepath),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    accent = colors.HexColor("#FF6B35")

    style_h1 = ParagraphStyle(
        "CustomH1",
        parent=styles["Title"],
        fontSize=18,
        textColor=accent,
        spaceAfter=8,
        alignment=1,
    )
    style_h2 = ParagraphStyle(
        "CustomH2",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#333333"),
        spaceBefore=10,
        spaceAfter=4,
        borderPadding=(0, 0, 2, 0),
    )
    style_h3 = ParagraphStyle(
        "CustomH3",
        parent=styles["Heading3"],
        fontSize=11,
        spaceBefore=6,
        spaceAfter=2,
    )
    style_body = ParagraphStyle(
        "CustomBody", parent=styles["Normal"], fontSize=10, spaceAfter=4, leading=14
    )
    style_bullet = ParagraphStyle(
        "CustomBullet",
        parent=styles["Normal"],
        fontSize=10,
        leftIndent=18,
        spaceAfter=2,
        bulletText="•",
    )

    story = []
    sections = _parse_markdown_to_sections(markdown_content)

    for stype, content in sections:
        if stype == "separator" or not content:
            story.append(Spacer(1, 6))
            continue
        if stype == "h1":
            story.append(Paragraph(content, style_h1))
            story.append(HRFlowable(width="100%", thickness=1, color=accent, spaceAfter=6))
        elif stype == "h2":
            story.append(Paragraph(content, style_h2))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc"), spaceAfter=4))
        elif stype == "h3":
            story.append(Paragraph(content, style_h3))
        elif stype == "bold":
            story.append(Paragraph(f"<b>{content}</b>", style_body))
        elif stype == "bullet":
            story.append(Paragraph(f"• {content}", style_bullet))
        else:
            story.append(Paragraph(content, style_body))

    doc.build(story)
    return filename


def get_file_path(filename: str) -> Path:
    return TEMP_DIR / filename
